import pdfplumber
import docx
from io import BytesIO

def parse_resume(file):
    filename = file.filename.lower()
    text = ""

    # -------- PDF Handling --------
    if filename.endswith(".pdf"):
        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        text += txt + "\n"
        except:
            pass  # continue to fallback

        # If text insufficient → fallback using PyPDF2
        if len(text.strip()) < 30:
            try:
                from PyPDF2 import PdfReader
                file.seek(0)
                reader = PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() or ""
            except:
                return ""
    
    # -------- DOCX Handling --------
    elif filename.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"

    # -------- TXT Handling --------
    elif filename.endswith(".txt"):
        text = file.read().decode("utf-8", errors="ignore")

    return text.strip()
import pdfplumber
import docx
from io import BytesIO

def parse_resume(file):
    filename = file.filename.lower()
    text = ""

    # -------- PDF Handling --------
    if filename.endswith(".pdf"):
        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        text += txt + "\n"
        except:
            pass  # continue to fallback

        # If text insufficient → fallback using PyPDF2
        if len(text.strip()) < 30:
            try:
                from PyPDF2 import PdfReader
                file.seek(0)
                reader = PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() or ""
            except:
                return ""
    
    # -------- DOCX Handling --------
    elif filename.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"

    # -------- TXT Handling --------
    elif filename.endswith(".txt"):
        text = file.read().decode("utf-8", errors="ignore")

    return text.strip()
